import re
import datetime
import pandas as pd
import logging
from typing import Optional, Dict, List, Tuple
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import streamlit as st
import fitz
import io
import base64
import requests

# Configurazione layout e stile Streamlit
st.set_page_config(layout="wide")

# Funzione per formattare i numeri
def format_number(value: float) -> str:
    """Format a number with dots as thousand separators and comma as decimal separator."""
    parts = "{:,.2f}".format(value).split('.')
    integer_part = parts[0].replace(',', '.')
    decimal_part = parts[1]
    return f"{integer_part},{decimal_part}"

# Funzione per normalizzare i nomi delle società
def normalizza_societa(nome_societa: str) -> str:
    if not nome_societa or nome_societa == "N/D":
        return nome_societa
    normalizzazione_map = {
        r'(?i)fiora(\s*s\.?p\.?a\.?)?$': 'ACQUEDOTTO DEL FIORA S.P.A.',
        r'(?i)acquedotto\s*del\s*fiora(\s*s\.?p\.?a\.?)?$': 'ACQUEDOTTO DEL FIORA S.P.A.',
        r'(?i)fiora\s*spa$': 'ACQUEDOTTO DEL FIORA S.P.A.',
        r'(?i)fiora\s*s\.p\.a\.$': 'ACQUEDOTTO DEL FIORA S.P.A.'
    }
    for pattern, replacement in normalizzazione_map.items():
        if re.search(pattern, nome_societa):
            return replacement
    return nome_societa

# Dizionario delle partite IVA delle società comuni
PIva_DATABASE = {
    "AGSM AIM ENERGIA S.P.A.": "01584620234",
    "A2A ENERGIA S.P.A.": "12883420155",
    "ACQUE VERONA S.P.A.": "02352230235",
    "ACQUE S.P.A.": "05006920482",
    "ACQUEDOTTO DEL FIORA S.P.A.": "01153850523",
    "ASA LIVORNO S.P.A.": "00102150497",
    "ENEL ENERGIA S.P.A.": "00934061007",
    "NUOVE ACQUE S.P.A.": "01359930482",
    "GAIA S.P.A.": "01966240465",
    "PUBLIACQUA S.P.A.": "01645330482",
    "EDISON ENERGIA S.P.A.": "09514811001",
    "G.E.A.L. S.P.A.": "01494020462",
    "Firenze Acqua SRL": "03671970485",
    "S.E.M.P. S.R.L.": "00281510453"
}

st.markdown("""
    <style>
        #MainMenu {visibility: hidden;}
        footer {visibility: hidden;}
        header {visibility: hidden;}
        .main .block-container {
            padding-top: 1rem;
            padding-right: 1rem;
            padding-left: 1rem;
            padding-bottom: 1rem;
        }
    </style>
""", unsafe_allow_html=True)

# Configurazione del logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Mappa mesi in italiano
MESI_MAP = {
    "gennaio": 1, "febbraio": 2, "marzo": 3, "aprile": 4, "maggio": 5, "giugno": 6,
    "luglio": 7, "agosto": 8, "settembre": 9, "ottobre": 10, "novembre": 11, "dicembre": 12
}

# Elenco esteso di società conosciute con regex specifiche
SOCIETA_CONOSCIUTE = {
    "AGSM AIM ENERGIA S.P.A.": r"AGSM\s*AIM\s*ENERGIA",
    "A2A ENERGIA S.P.A.": r"A2A\s*ENERGIA",
    "ACQUE VERONA S.P.A.": r"ACQUE\s*VERONA",
    "ACQUE S.P.A.": r"ACQUE\s*S\.?P\.?A\.?",
    "ACQUEDOTTO DEL FIORA S.P.A.": r"ACQUEDOTTO\s*DEL\s*FIORA|FIORA\s*S\.?P\.?A\.?",
    "ASA LIVORNO S.P.A.": r"ASA\s*LIVORNO",
    "ENEL ENERGIA S.P.A.": r"ENEL\s*ENERGIA",
    "NUOVE ACQUE S.P.A.": r"NUOVE\s*ACQUE",
    "GAIA S.P.A.": r"GAIA\s*S\.?P\.?A\.?",
    "PUBLIACQUA S.P.A.": r"PUBLIACQUA",
    "EDISON ENERGIA S.P.A.": r"EDISON\s*ENERGIA",
    "G.E.A.L. S.P.A.": r"G\.?E\.?A\.?L\.?\s*S\.?P\.?A\.?",
    "Firenze Acqua SRL": r"FIRENZE\s*ACQUA\s*S\.?R\.?L\.?",
    "S.E.M.P. S.R.L.": r"S\.?E\.?M\.?P\.?\s*S\.?R\.?L\.?"
}

def estrai_testo_da_pdf(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        testo = ""
        for page in doc:
            testo += page.get_text()
        return testo
    except fitz.FileDataError:
        logger.error(f"File {file.name} non valido o corrotto")
        return ""
    except Exception as e:
        logger.error(f"Errore durante l'estrazione del testo dal PDF {file.name}: {str(e)}")
        return ""

def estrai_societa(testo: str) -> str:
    try:
        for societa, pattern in SOCIETA_CONOSCIUTE.items():
            if re.search(pattern, testo, re.IGNORECASE):
                return normalizza_societa(societa)
        patterns = [
            r'\b([A-Z]{2,}\s*(?:AIM|ENERGIA|GAS|ACQUA|SPA))\b',
            r'\b(SPA|S\.P\.A\.|SRL|S\.R\.L\.)\b'
        ]
        for pattern in patterns:
            match = re.search(pattern, testo)
            if match:
                return normalizza_societa(match.group(0).strip())
    except Exception as e:
        logger.error(f"Errore durante l'estrazione della società: {str(e)}")
    return "N/D"

def estrai_periodo(testo: str) -> str:
    try:
        patterns = [
            r'dal\s+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\s+al\s+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'periodo\s+di\s+riferimento\s*:\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\s*-\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'Periodo di riferimento\s*:\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\s*-\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'rif\.\s*periodo\s*(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})\s+al\s+(\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4})',
            r'dal\s+(\d{1,2}/\d{1,2}/\d{4})\s+al\s+(\d{1,2}/\d{1,2}/\d{4})',
            r'Periodo di riferimento\s+(\d{1,2}/\d{1,2}/\d{4}\s*-\s*\d{1,2}/\d{1,2}/\d{4})',
            r'Periodo\s*:\s*(\d{1,2}/\d{1,2}/\d{4})\s*-\s*(\d{1,2}/\d{1,2}/\d{4})',
            r'Periodo fatturazione\s*:\s*(\d{1,2}/\d{1,2}/\d{4})\s*-\s*(\d{1,2}/\d{1,2}/\d{4})',
            r'dal\s+(\d{2}/\d{2}/\d{4})\s+al\s+(\d{2}/\d{2}/\d{4})',
            r'Periodo di riferimento (\d{2}/\d{2}/\d{4}) - (\d{2}/\d{2}/\d{4})',
            r'(\d{2}/\d{2}/\d{4}) - (\d{2}/\d{2}/\d{4})'
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, testo, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) == 2:
                    return f"{match.group(1)} - {match.group(2)}"
    except Exception as e:
        logger.error(f"Errore durante l'estrazione del periodo: {str(e)}")
    return "N/D"

def parse_date(g: str, m: str, y: str) -> Optional[datetime.date]:
    try:
        giorno = int(g)
        if m.isdigit():
            mese = int(m)
        else:
            mese = MESI_MAP.get(m.lower().strip(), 0)
        if len(y) == 2:
            anno = 2000 + int(y)
        else:
            anno = int(y)
        if 1 <= mese <= 12 and 1 <= giorno <= 31:
            return datetime.date(anno, mese, giorno)
    except (ValueError, TypeError) as e:
        logger.error(f"Errore durante il parsing della data: {str(e)}")
    return None

def estrai_data_fattura(testo: str) -> str:
    try:
        patterns = [
            r'(?:data\s*fattura|fattura\s*del|emissione)\s*[:\-]?\s*(\d{1,2})[\/\-\.\s](\d{1,2}|\w+)[\/\-\.\s](\d{2,4})',
            r'Bolletta\s*n\.\s*\d+\s*del\s*(\d{1,2})\s*(\w+)\s*(\d{4})',
            r'(?:data\s*emissione|emesso\s*il)\s*[:\-]?\s*(\d{1,2})[\/\-\.\s](\d{1,2}|\w+)[\/\-\.\s](\d{2,4})',
            r'\b(\d{2})[\/\-\.](\d{2})[\/\-\.](\d{4})\b',
            r'\b(\d{4})[\/\-\.](\d{2})[\/\-\.](\d{2})\b',
            r'\b(\d{1,2})\s+(gennaio|febbraio|marzo|aprile|maggio|giugno|luglio|agosto|settembre|ottobre|novembre|dicembre)\s+(\d{4})\b',
            r'\b(?:al|il)\s+(\d{1,2})\s+(\w+)\s+(\d{4})\b'
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, testo, re.IGNORECASE)
            for match in matches:
                if len(match.groups()) == 3:
                    data = parse_date(match.group(1), match.group(2), match.group(3))
                    if data:
                        return data.strftime("%d/%m/%Y")
    except Exception as e:
        logger.error(f"Errore durante l'estrazione della data: {str(e)}")
    return "N/D"

def estrai_pod_pdr(testo: str) -> str:
    try:
        pod_patterns = [
            r'POD\s*[:\-]?\s*([A-Z0-9]{14,16})',
            r'Punto\s*di\s*Prelievo\s*[:\-]?\s*([A-Z0-9]{14,16})',
            r'Codice\s*POD\s*[:\-]?\s*([A-Z0-9]{14,16})',
            r'(?:matricola\s*contatore|matr\.?\s*cont\.?|numero\s*contatore)\s*[:=\-]?\s*([A-Z0-9]{8,12})(?:\s|$)',
            r'(?:matricola\s*contatore|matr\.?\s*cont\.?|numero\s*contatore)\s*[:=\-]?\s*([A-Z0-9\-]{8,12})(?:\s|$)',
            r'(?:matricola\s*contatore|matr\.?\s*cont\.?|numero\s*contatore)\s*[:=\-]?\s*([A-Z0-9]{8,14})(?:\s|$)',
            r'Contatore\s*n\.\s*(\d{6,})',
            r'Matricola\s*Misuratore\s*:\s*(\d+)'
        ]
        for pattern in pod_patterns:
            match = re.search(pattern, testo, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        pdr_patterns = [
            r'PDR\s*[:\-]?\s*([A-Z0-9]{14,16})',
            r'Punto\s*di\s*Ricerca\s*[:\-]?\s*([A-Z0-9]{14,16})',
            r'Codice\s*PDR\s*[:\-]?\s*([A-Z0-9]{14,16})'
        ]
        for pattern in pdr_patterns:
            match = re.search(pattern, testo, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    except Exception as e:
        logger.error(f"Errore durante l'estrazione del POD/PDR: {str(e)}")
    return "N/D"

import re


def estrai_indirizzo(testo: str) -> str:
    try:

        # Pattern per cercare all'interno del riquadro "DATI DELLA FORNITURA"
        pattern_dati_fornitura = r'DATI DELLA FORNITURA(.*?)(?=\n\n|\Z)'
        match_dati_fornitura = re.search(pattern_dati_fornitura, testo, re.IGNORECASE | re.DOTALL)

        if match_dati_fornitura:
            sezione_fornitura = match_dati_fornitura.group(1)

            # Pattern per cercare l'indirizzo all'interno della sezione trovata
            pattern_indirizzo = r'(?:Indirizzo di fornitura|VIA|Viale|Piazza|Corso)\s*:\s*([^\n|e-mail]+)'
            match_indirizzo = re.search(pattern_indirizzo, sezione_fornitura, re.IGNORECASE)

            if match_indirizzo:
                indirizzo = match_indirizzo.group(1).strip()
                indirizzo = re.sub(r'^\W+|\W+$', '', indirizzo)
                return indirizzo
 
        # Pattern per Nuove Acque
        pattern_nuove_acque = r'Indirizzo\s+fornitura\s+([^\n]+)\s*-\s*\d{5}\s+[A-Z]{2}'
        match_nuove_acque = re.search(pattern_nuove_acque, testo, re.IGNORECASE)
        if match_nuove_acque:
            return match_nuove_acque.group(1).strip()

        # Pattern per Gaia
        pattern_gaia = r'INTESTAZIONE\s*([^\n]+)\s*([^\n]+)\s*(\d{5}\s+[A-Z]{2})'
        match_gaia = re.search(pattern_gaia, testo, re.IGNORECASE | re.DOTALL)
        if match_gaia:
            return match_gaia.group(2).strip()

        # Pattern per Fiora
        pattern_fiora = (
            r'(?:DATI FORNITURA|Indirizzo[^\n]*)\s*'
            r'(?:.*\n)*?'
            r'((?:VIA|CORSO|PIAZZA|STRADA|V\.|C\.SO|P\.ZA)\s?.+?\d{1,5}(?:\s*[A-Za-z]?)?)\b'
        )
        match_fiora = re.search(pattern_fiora, testo, re.IGNORECASE | re.MULTILINE)
        if match_fiora:
            indirizzo = match_fiora.group(1).strip()
            indirizzo = re.sub(r'^\W+|\W+$', '', indirizzo)
            return indirizzo

        # Pattern generici
        patterns_generici = [
            r'Indirizzo\s*[:\-]?\s*((?:Via|Viale|Piazza|Corso|C\.so|C\.|V\.le|Str\.|C.so|V\.|P\.za).+?\d{1,5}(?:\s*[A-Za-z]?)?)\b',
            r'Servizio\s*erogato\s*in\s*((?:Via|Viale|Piazza|Corso|C\.so|C\.|V\.le|Str\.|C.so|V\.|P\.za).+?\d{1,5}(?:\s*[A-Za-z]?)?)\b',
            r'Luogo\s*di\s*fornitura\s*[:\-]?\s*((?:Via|Viale|Piazza|Corso|C\.so|C\.|V\.le|Str\.|C.so|V\.|P\.za).+?\d{1,5}(?:\s*[A-Za-z]?)?)\b',
            r'Indirizzo\s*di\s*fornitura\s*[:\-]?\s*((?:Via|Viale|Piazza|Corso|C\.so|C\.|V\.le|Str\.|C.so|V\.|P\.za).+?\d{1,5}(?:\s*[A-Za-z]?)?)\b',
            r'Indirizzo\s*fornitura\s*((?:Via|Viale|Piazza|Corso|C\.so|C\.|V\.le|Str\.|C.so|V\.|P\.za).+?\d{1,5}(?:\s*[A-Za-z]?)?)\b',
        ]

        for pattern in patterns_generici:
            match = re.search(pattern, testo, re.IGNORECASE | re.DOTALL)
            if match:
                indirizzo = match.group(1).strip()
                indirizzo = re.sub(r'^\W+|\W+$', '', indirizzo)
                indirizzo = re.sub(r'\s+', ' ', indirizzo)
                return indirizzo

        return "N/D"
    except Exception as e:
        print(f"Errore durante l'estrazione dell'indirizzo: {str(e)}")
        return "N/D"


def estrai_numero_fattura(testo: str) -> str:
    try:
        patterns = [
            r'Numero fattura elettronica valida ai fini fiscali\s*[:]?\s*([A-Z]{0,4}\s*[0-9\/\-]+\s*[0-9]+)',
            r'(\d{12})\s*numero\s*fattura\s*elettronica\s*valido\s*ai\s*fini\s*fiscali',
            r'(?:numero\s*fattura|n°\s*fattura|fattura\s*n\.?)\s*[:\-]?\s*([A-Z]{0,4}\s*[0-9\/\-]+\s*[0-9]+)',
            r'(?:doc\.|documento)\s*[:\-]?\s*([A-Z]{0,4}\s*[0-9\/\-]+\s*[0-9]+)',
            r'[Ff]attura\s+(?:elektronica\s+)?[nN]°?\s*[:\-]?\s*([A-Z]{0,4}\s*[0-9\/\-]+\s*[0-9]+)',
            r'Numero Fattura\s*[:]?\s*([A-Z]{0,4}\s*[0-9\/\-]+\s*[0-9]+)',
            r'\b\d{2,4}[\/\-]\d{3,8}\b',
            r'\b[A-Z]{2,5}\s*\d{4,}\/\d{2,}\b'
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, testo, re.IGNORECASE)
            for match in matches:
                num = match.group(1) if match.groups() else match.group(0)
                num = num.strip()
                if len(num) >= 5 and any(c.isdigit() for c in num):
                    return num
    except Exception as e:
        logger.error(f"Errore durante l'estrazione del numero della fattura: {str(e)}")
    return "N/D"

def estrai_totale_bolletta(testo: str) -> Tuple[str, str]:
    try:
        patterns = [
            r'totale\s*(?:fattura|bolletta)\s*[:\-]?\s*[€]?\s*([\d\.,]+)\s*([€]?)',
            r'importo\s*totale\s*[:\-]?\s*[€]?\s*([\d\.,]+)\s*([€]?)',
            r'pagare\s*[:\-]?\s*[€]?\s*([\d\.,]+)\s*([€]?)',
            r'totale\s*dovuto\s*[:\-]?\s*[€]?\s*([\d\.,]+)\s*([€]?)',
            r'TOTALE\s+Scissione\s+dei\s+pagamenti\s*[:\-]?\s*[€]?\s*([\d\.,]+)\s*([€]?)'
        ]
        for pattern in patterns:
            match = re.search(pattern, testo, re.IGNORECASE)
            if match and len(match.groups()) >= 1:
                importo = match.group(1).replace('.', '').replace(',', '.')
                try:
                    importo_float = float(importo)
                    valuta = match.group(2) if len(match.groups()) >= 2 and match.group(2) else "€"
                    return importo, valuta
                except ValueError:
                    continue
    except Exception as e:
        logger.error(f"Errore durante l'estrazione del totale della bolletta: {str(e)}")
    return "N/D", "€"

def determina_tipo_bolletta(societa: str, testo: str) -> str:
    societa_lower = societa.lower()
    testo_lower = testo.lower()
    if "agsm" in societa_lower:
        if "gas" in testo_lower:
            return "gas"
        else:
            return "energia"
    if any(kw in societa_lower for kw in ["acqua", "acquedotto", "fiora", "nuove acque", "pubbliacqua", "gaia", "acque", "asa"]):
        return "acqua"
    elif any(kw in societa_lower for kw in ["energia", "enel", "a2a", "edison"]):
        return "energia"
    elif any(kw in societa_lower for kw in ["gas"]):
        return "gas"
    else:
        return "sconosciuto"

def estrai_consumi(testo: str, tipo_bolletta: str) -> str:
    try:
        testo_upper = testo.upper()
        idx = testo_upper.find("RIEPILOGO CONSUMI FATTURATI")
        if idx != -1:
            snippet = testo_upper[idx:idx+600]
            match = re.search(r'TOTALE COMPLESSIVO DI[:\-]?\s*([\d\.,]+)', snippet)
            if not match:
                match = re.search(r'TOTALE\s+QUANTITÀ[:\-]?\s*([\d\.,]+)', snippet)
            if match:
                try:
                    valore = float(match.group(1).replace('.', '').replace(',', '.'))
                    if tipo_bolletta == "acqua":
                        return f"{valore} mc"
                    elif tipo_bolletta == "energia":
                        return f"{valore} kWh"
                    elif tipo_bolletta == "gas":
                        return f"{valore} Smc"
                except:
                    pass
        patterns = [
            r'consumo\s*([\d\.]+)\s*kWh',
            r'Consumo\s*\n\s*(\d+)\s*mc',
            r'Consumo\s+nel\s+periodo\s+di\s+\d+\s+giorni:\s*([\d\.,]+)\s*mc',
            r'Letture e Consumi.*?Contatore n\.\s*\d+.*?(\d+)\s*mc',
            r'Consumo\s+stimato\s*[:\-]?\s*([\d\.,]+)\s*mc',
            r'Consumo\s+fatturato\s*[:\-]?\s*([\d\.,]+)\s*mc',
            r'totale\s+smc\s+fatturati\s*[:\-]?\s*([\d]{1,3}(?:[\.,][\d]{3})*(?:[\.,]\d+)?)',
            r'Totale\s+quantità\s*[:\-]?\s*([\d.]+,\d+)\s*Smc',
            r'totale\s+consumo\s+fatturato\s+per\s+il\s+periodo\s+di\s+riferimento\s*[:\-]?\s*([\d\.,]+)\s*(mc|m³|metri\s*cubi)',
            r'(?:consumo\s*fatturato|consumo\s*stimato\s*fatturato|consumo\s*totale)\s*[:\-]?\s*([\d\.,]+)\s*(mc|m³|metri\s*cubi)',
            r'(?:riepilogo\s*consumi[^\n]*\n.*\n.*?)([\d\.,]+)\s*(mc|m³|metri\s*cubi)',
            r'(?:prospetto\s*letture\s*e\s*consumi[^\n]*\n.*\n.*?\d+)\s+([\d\.,]+)\s*$',
            r'(?:dettaglio\s*consumi[^\n]*\n.*\n.*?\d+\s+)([\d\.,]+)\s*$',
            r'Consumo\s+([\d\.]+)\s*mc',
            r'Consumo\s+del\s+periodo:\s*([\d\.,]+)\s*mc',
            r'Consumo\s+fatturato\s*[:\-]?\s*([\d\.,]+)\s*mc',
            r'Consumo\s+stimato\s*[:\-]?\s*([\d\.,]+)\s*mc',
            r'Consumo\s+effettivo\s*[:\-]?\s*([\d\.,]+)\s*mc',
            r'Consumo\s+([\d\.]+)\s*metri\s*cubi',
            r'Consumo\s+([\d\.]+)\s*m³'
        ]
        for pattern in patterns:
            matches = re.finditer(pattern, testo, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                try:
                    valore_raw = match.group(1)
                    valore_normalizzato = valore_raw.replace('.', '').replace(',', '.')
                    consumo = float(valore_normalizzato)
                    if len(match.groups()) > 1 and match.group(2):
                        unita = match.group(2).lower()
                    else:
                        if tipo_bolletta == "acqua":
                            unita = "mc"
                        elif tipo_bolletta == "energia":
                            unita = "kWh"
                        elif tipo_bolletta == "gas":
                            unita = "Smc"
                        else:
                            unita = "mc"
                    return f"{consumo} {unita}"
                except (ValueError, IndexError):
                    continue
        fallback = re.search(r'(\d+)\s*mc\s+Importo\s+da\s+pagare', testo)
        if fallback:
            return f"{float(fallback.group(1))} mc"
    except Exception as e:
        logger.error(f"Errore durante l'estrazione dei consumi: {str(e)}", exc_info=True)
    return "N/D"

def estrai_dati_cliente(testo: str) -> str:
    try:
        patterns = [
            r'(?:Numero\s*Contatore|Contatore)[\s:]*([0-9]{8,9})',
            r'(?:Matricola|Contatore|S/N)[\s:]*([A-Z0-9]{14,15})'
        ]
        for pattern in patterns:
            match = re.search(pattern, testo, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return "N/D"
    except Exception as e:
        logger.error(f"Errore durante l'estrazione dei dati cliente: {str(e)}")
        return "N/D"

def estrai_dati(file):
    testo = estrai_testo_da_pdf(file)
    if not testo:
        return None
    societa = estrai_societa(testo)
    tipo_bolletta = determina_tipo_bolletta(societa, testo)
    pod = estrai_pod_pdr(testo)
    totale, valuta = estrai_totale_bolletta(testo)
    consumi = estrai_consumi(testo, tipo_bolletta)
    indirizzo = estrai_indirizzo(testo)
    dati_cliente = estrai_dati_cliente(testo)
    return {
        "Società": societa,
        "Periodo di Riferimento": estrai_periodo(testo),
        "Data Fattura": estrai_data_fattura(testo),
        "POD": pod,
        "Dati Cliente": dati_cliente,
        "Indirizzo": indirizzo,
        "Numero Fattura": estrai_numero_fattura(testo),
        f"Totale ({valuta})": format_number(float(totale.replace(',', '.'))) if totale != "N/D" else totale,
        "File": file.name,
        "Consumi": consumi
    }

def crea_excel(dati_lista: List[Dict[str, str]]):
    try:
        colonne_ordinate = [
            "Società",
            "Periodo di Riferimento",
            "Data Fattura",
            "POD",
            "Dati Cliente",
            "Indirizzo",
            "Numero Fattura",
            "Totale (€)",
            "File",
            "Consumi"
            
        ]
        df = pd.DataFrame([d for d in dati_lista if d is not None])
        if len(df) == 0:
            st.warning("Nessun dato valido da esportare")
            return None
        colonne_presenti = [col for col in colonne_ordinate if col in df.columns]
        df = df[colonne_presenti]
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df.to_excel(writer, index=False, sheet_name='Report')
            workbook = writer.book
            worksheet = writer.sheets['Report']
            header_format = workbook.add_format({
                'bold': True,
                'text_wrap': True,
                'valign': 'top',
                'fg_color': '#4472C4',
                'font_color': 'white',
                'border': 1
            })
            data_format = workbook.add_format({
                'text_wrap': True,
                'valign': 'top',
                'border': 1
            })
            for col_num, value in enumerate(df.columns.values):
                worksheet.write(0, col_num, value, header_format)
            for row in range(1, len(df)+1):
                for col in range(len(df.columns)):
                    worksheet.write(row, col, df.iloc[row-1, col], data_format)
            for i, col in enumerate(df.columns):
                max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
                worksheet.set_column(i, i, max_len)
        output.seek(0)
        return output
    except Exception as e:
        logger.error(f"Errore durante la creazione del file Excel: {str(e)}")
        return None

def mostra_grafico_consumi(dati_lista: List[Dict[str, str]]):
    try:
        df = pd.DataFrame([d for d in dati_lista if d is not None])
        if len(df) == 0:
            return
        if "Consumi" not in df.columns:
            return
        df['Consumo_val'] = df['Consumi'].str.extract(r'([\d\.]+)')[0].astype(float)
        df = df.dropna(subset=['Consumo_val'])
        if len(df) < 2:
            return
        st.subheader("📈 Confronto Consumi")
        unita = df['Consumi'].iloc[0].split()[-1] if len(df['Consumi'].iloc[0].split()) > 1 else ""
        chart_data = df[['File', 'Consumo_val']].rename(columns={'Consumo_val': 'Consumo'})
        chart_data = chart_data.set_index('File')
        st.bar_chart(chart_data)
        if unita:
            st.caption(f"Unità di misura: {unita}")
    except Exception as e:
        st.warning(f"Impossibile generare il grafico: {str(e)}")

def crea_attestazione(dati: List[Dict[str, str]], firma_selezionata: str = "Mar. Basile Vincenzo"):
    try:
        doc = Document()
        section = doc.sections[0]

        # Ridurre ulteriormente i margini laterali
        section.left_margin = Pt(80)  # Ridotto ulteriormente
        section.right_margin = Pt(80)  # Ridotto ulteriormente

        section.top_margin = Pt(50)
        section.bottom_margin = Pt(50)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        data_fattura_str = dati[0].get('Data Fattura') if dati else None
        if not data_fattura_str:
            raise ValueError("Data fattura non presente nei dati")
        try:
            data_fattura = datetime.datetime.strptime(data_fattura_str, "%d/%m/%Y")
        except ValueError:
            raise ValueError(f"Formato data fattura non valido: {data_fattura_str}. Atteso GG/MM/AAAA")

        if data_fattura.weekday() == 5:  # Sabato
            data_attestazione = data_fattura - datetime.timedelta(days=1)
        elif data_fattura.weekday() == 6:  # Domenica
            data_attestazione = data_fattura - datetime.timedelta(days=2)
        else:
            data_attestazione = data_fattura

        logo_url = "https://upload.wikimedia.org/wikipedia/commons/thumb/0/00/Emblem_of_Italy.svg/1200px-Emblem_of_Italy.svg.png"
        try:
            header = doc.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            response = requests.get(logo_url)
            if response.status_code == 200:
                logo_stream = io.BytesIO(response.content)
                header.add_run().add_picture(logo_stream, width=Pt(56.5), height=Pt(56.5))
            header.add_run("\n\n")
            header_run = header.add_run("Guardia di Finanza\n")
            header_run.bold = True
            header_run.font.size = Pt(20)
            header_run.font.name = 'Arial'
            header_run = header.add_run("REPARTO TECNICO LOGISTICO AMMINISTRATIVO TOSCANA\n")
            header_run.bold = True
            header_run.font.size = Pt(16)
            header_run.font.name = 'Arial'
            header_run = header.add_run("Ufficio Logistico - Sezione Infrastrutture\n\n")
            header_run.bold = True
            header_run.font.size = Pt(14)
            header_run.font.name = 'Arial'
        except Exception as e:
            logger.error(f"Errore durante l'aggiunta del logo: {str(e)}")
            header_run = header.add_run("Guardia di Finanza\n")
            header_run.bold = True
            header_run.font.size = Pt(20)
            header_run.font.name = 'Arial'
            header_run = header.add_run("REPARTO TECNICO LOGISTICO AMMINISTRATIVO TOSCANA\n")
            header_run.bold = True
            header_run.font.size = Pt(16)
            header_run.font.name = 'Arial'
            header_run = header.add_run("Ufficio Logistico - Sezione Infrastrutture\n\n")
            header_run.bold = True
            header_run.font.size = Pt(14)
            header_run.font.name = 'Arial'

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_format = title.paragraph_format
        title_format.space_before = Pt(0)
        title_format.space_after = Pt(0)
        title_format.border_top = Pt(1)
        title_format.border_bottom = Pt(1)
        title_format.border_left = Pt(1)
        title_format.border_right = Pt(1)
        title_format.border_top_color = RGBColor(0, 0, 0)
        title_format.border_bottom_color = RGBColor(0, 0, 0)
        title_format.border_left_color = RGBColor(0, 0, 0)
        title_format.border_right_color = RGBColor(0, 0, 0)
        title_format.space_inside = Pt(4)
        title_run = title.add_run("Dichiarazione di regolare fornitura")
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_run.font.name = 'Arial'

        societa = normalizza_societa(dati[0].get('Società', 'ACQUE S.P.A.')) if dati else 'ACQUE S.P.A.'
        tipo_fornitura = determina_tipo_bolletta(societa, "")

        body_text = (
            "Si attesta l'avvenuta attività di controllo tecnico-logistica come da circolare "
            "90000/310 edizione 2011 del Comando Generale G. di F. - I Reparto Ufficio Ordinamento - "
            "aggiornata con circolare nr. 209867/310 del 06.07.2016.\n\n"
            "Si dichiara che i costi riportati nelle seguenti fatture elettroniche:\n"
        )
        body = doc.add_paragraph(body_text)
        body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'N. Documento'
        hdr_cells[1].text = 'Data Fattura'
        hdr_cells[2].text = 'Totale (€)'

        for fattura in dati:
            row_cells = table.add_row().cells
            row_cells[0].text = fattura.get('Numero Fattura', 'N/D')
            row_cells[1].text = fattura.get('Data Fattura', 'N/D')
            row_cells[2].text = fattura.get('Totale (€)', 'N/D')

        for i, cell in enumerate(table.columns):
            max_length = max(len(str(row.cells[i].text)) for row in table.rows)
            for row in table.rows:
                row.cells[i].width = Pt(max_length * 10)

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.alignment = 1

        # Ridurre lo spazio dopo la tabella
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

        piva = dati[0].get('P.IVA')
        if not piva:
            piva = PIva_DATABASE.get(societa)
            if not piva:
                piva = PIva_DATABASE["ACQUE S.P.A."]
                logger.warning(f"P.IVA non trovata per società: {societa}. Usato valore default ACQUE S.P.A.")

        if societa == "A2A ENERGIA S.P.A.":
            footer_text = (
                "emessa dalla società A2A ENERGIA S.P.A. - P.I. {} - "
                "nell'ambito della convenzione CONSIP \"Fornitura Energia Elettrica 12 Mesi - Lotto 8 Toscana\" "
                "(Codice Identificativo Gara: B349419163), si riferiscono effettivamente a consumi di energia elettrica "
                "effettuati dai Comandi amministrati da questo Reparto per i fini istituzionali.\n\n"
                "L'energia elettrica oggetto della prefata fattura è stata regolarmente erogata "
                "presso i contatori richiesti dall'Amministrazione, ubicati presso le caserme del Corpo dislocate nella Regione Toscana.\n".format(piva)
            )
        else:
            if tipo_fornitura == "acqua":
                footer_text = (
                    "\nemesse dalla società {} -- P.I. {} -- si riferiscono effettivamente a "
                    "consumi di acqua effettuati dai Comandi amministrati da questo Reparto per i fini istituzionali.\n\n"
                    "L'acqua oggetto delle prefate fatture è stata regolarmente erogata presso i contatori richiesti "
                    "dall'Amministrazione, ubicati presso le caserme del Corpo dislocate nella Regione Toscana.\n".format(societa, piva)
                )
            else:
                footer_text = (
                    "\nemesse dalla società {} -- P.I. {} -- si riferiscono effettivamente a "
                    "consumi di materia prima effettuati dai Comandi amministrati da questo Reparto per i fini istituzionali.\n\n"
                    "La materia prima oggetto delle prefate fatture è stata regolarmente erogata presso i contatori richiesti "
                    "dall'Amministrazione, ubicati presso le caserme del Corpo dislocate nella Regione Toscana.\n".format(societa, piva)
                )

        footer = doc.add_paragraph(footer_text)
        footer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        data_attestazione_str = data_attestazione.strftime("%d.%m.%Y")
        data_para = doc.add_paragraph(f"\nFirenze, {data_attestazione_str}\n")
        data_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Migliorare il gruppo firma incolonnandolo
        firma_paragraph = doc.add_paragraph()
        firma_run = firma_paragraph.add_run("L'Addetto al Drappello Gestione Patrimonio Immobiliare")
        firma_run.font.name = 'Arial'
        firma_run.font.size = Pt(12)
        firma_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        firma_paragraph = doc.add_paragraph()
        firma_run = firma_paragraph.add_run(firma_selezionata)
        firma_run.font.name = 'Arial'
        firma_run.font.size = Pt(12)
        firma_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        nome_societa_pulito = re.sub(r'[^a-zA-Z0-9]', '_', societa)
        nome_file = f"attestazione_{nome_societa_pulito}_{data_attestazione.strftime('%Y%m%d')}.docx"
        return output, nome_file
    except Exception as e:
        logger.error(f"Errore durante la creazione dell'attestazione: {str(e)}")
        return None, "attestazione.docx"

def main():
    st.title("📊 REPORT 2.0")
    st.markdown("**Carica una o più bollette PDF** per estrarre automaticamente i dati principali.")
    st.markdown("""
    <style>
    div[data-baseweb="base-input"] {
        width: 100%;
    }
    div[data-testid="stDataFrame"] {
        width: 100%;
    }
    </style>
    """, unsafe_allow_html=True)
    with st.sidebar:
        st.header("Impostazioni")
        mostra_grafici = st.checkbox("Mostra grafici comparativi", value=True)
        raggruppa_societa = st.checkbox("Raggruppa per società", value=True)
    file_pdf_list = st.file_uploader(
        "Seleziona i file PDF delle bollette",
        type=["pdf"],
        accept_multiple_files=True,
        help="Puoi selezionare più file contemporaneamente"
    )
    if file_pdf_list:
        risultati = []
        progress_bar = st.progress(0)
        status_text = st.empty()
        for i, file in enumerate(file_pdf_list):
            status_text.text(f"Elaborazione {i+1}/{len(file_pdf_list)}: {file.name[:30]}...")
            progress_bar.progress((i + 1) / len(file_pdf_list))
            try:
                dati = estrai_dati(file)
                if dati:
                    risultati.append(dati)
            except Exception as e:
                logger.error(f"Errore durante l'elaborazione di {file.name}: {str(e)}")
                continue
        progress_bar.empty()
        if risultati:
            status_text.success(f"✅ Elaborazione completata! {len(risultati)} file processati con successo.")
            st.subheader("📋 Dati Estratti")
            if raggruppa_societa:
                societa_disponibili = sorted(list(set(d['Società'] for d in risultati if pd.notna(d['Società']) and (d['Società'] != "N/D"))))
                if societa_disponibili:
                    societa = st.selectbox(
                        "Filtra per società",
                        options=["Tutte"] + societa_disponibili,
                        index=0
                    )
                    if societa != "Tutte":
                        risultati_filtrati = [d for d in risultati if d['Società'] == societa]
                    else:
                        risultati_filtrati = risultati
                else:
                    risultati_filtrati = risultati
                    st.warning("Nessuna società riconosciuta nei documenti")
            else:
                risultati_filtrati = risultati
            df = pd.DataFrame(risultati_filtrati)
            st.data_editor(
                df,
                use_container_width=True,
                hide_index=True,
                disabled=True,
                key="data_editor"
            )
            if mostra_grafici and risultati_filtrati:
                mostra_grafico_consumi(risultati_filtrati)
            st.subheader("📤 Esporta Dati")
            col1, col2, col3 = st.columns(3)
            with col1:
                excel_data = crea_excel(risultati_filtrati)
                if excel_data:
                    st.download_button(
                        label="Scarica Excel",
                        data=excel_data,
                        file_name="report_consumi.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        help="Scarica i dati in formato Excel"
                    )
            with col2:
                if risultati_filtrati:
                    csv = pd.DataFrame(risultati_filtrati).to_csv(index=False, sep=';').encode('utf-8')
                    st.download_button(
                        label="Scarica CSV",
                        data=csv,
                        file_name="report_consumi.csv",
                        mime="text/csv",
                        help="Scarica i dati in formato CSV (delimitato da punto e virgola)"
                    )
            with col3:
                if risultati_filtrati:
                    st.markdown("**Seleziona firma:**")
                    firma_selezionata = st.radio(
                        "Firma attestazione",
                        options=[
                            "Mar. Basile Vincenzo",
                            "Cap. Carla Mottola"
                        ],
                        index=0,
                        label_visibility="collapsed"
                    )
                    attestazione, nome_file = crea_attestazione(risultati_filtrati, firma_selezionata)
                    if attestazione:
                        st.download_button(
                            label="Scarica Attestazione",
                            data=attestazione,
                            file_name=nome_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            help="Scarica l'attestazione precompilata in formato Word"
                        )
                    else:
                        st.warning("Errore nella generazione dell'attestazione")
        else:
            status_text.warning("⚠️ Nessun dato valido estratto dai file caricati")
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; font-size: 14px; color: gray;">
        Strumento sviluppato dal Mar. Vincenzo Basile<br>
        Supporta i principali fornitori italiani di luce, gas e acqua
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
